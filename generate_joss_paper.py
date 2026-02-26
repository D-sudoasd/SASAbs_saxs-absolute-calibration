#!/usr/bin/env python3
"""Generate the JOSS paper draft as a Word document (.docx).

Run:  python generate_joss_paper.py
Output: paper/saxsabs_joss_paper.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── helpers ────────────────────────────────────────────────────────
PAPER_DIR = Path(__file__).resolve().parent / "paper"

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text, bold=False, italic=False, font_size=Pt(11)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = font_size
    run.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    return p

def add_rich_para(doc, segments):
    """segments: list of (text, bold, italic)"""
    p = doc.add_paragraph()
    for text, bold, italic in segments:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(6)
    return p

def add_equation(doc, text, label=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = "Cambria Math"
    if label:
        run2 = p.add_run(f"    ({label})")
        run2.font.size = Pt(11)
        run2.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    return p

def add_figure(doc, image_filename, caption, width=Inches(5.8)):
    """Insert an image with a centred caption below it."""
    img_path = PAPER_DIR / image_filename
    if not img_path.exists():
        add_para(doc, f"[Figure placeholder: {image_filename} not found]", italic=True)
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(str(img_path), width=width)
    p_img.paragraph_format.space_after = Pt(2)
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_cap.add_run(caption)
    run_cap.italic = True
    run_cap.font.size = Pt(9)
    run_cap.font.name = "Times New Roman"
    p_cap.paragraph_format.space_after = Pt(10)
    return p_img

# ── main ───────────────────────────────────────────────────────────
def build_paper():
    doc = Document()

    # -- Default style --
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # ================================================================
    # TITLE
    # ================================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(
        "saxsabs: A Robust Workflow for Small-Angle X-ray Scattering "
        "Absolute Intensity Calibration"
    )
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"
    title_p.paragraph_format.space_after = Pt(4)

    # Author
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author_p.add_run("Delun Gong")
    ar.font.size = Pt(12)
    ar.font.name = "Times New Roman"
    author_p.paragraph_format.space_after = Pt(2)

    # Affiliation
    aff_p = doc.add_paragraph()
    aff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    af = aff_p.add_run(
        "Institute of Metal Research, Chinese Academy of Sciences, Shenyang 110016, China"
    )
    af.font.size = Pt(10)
    af.italic = True
    af.font.name = "Times New Roman"
    aff_p.paragraph_format.space_after = Pt(12)

    # ================================================================
    # SUMMARY
    # ================================================================
    add_heading(doc, "Summary", level=1)

    add_para(doc,
        "saxsabs is an open-source Python package that provides a complete, "
        "reproducible workflow for small-angle X-ray scattering (SAXS) absolute "
        "intensity calibration. It automates the data-reduction chain from raw "
        "two-dimensional (2D) detector images to calibrated one-dimensional (1D) "
        "scattering profiles on an absolute cross-section scale (cm⁻¹ sr⁻¹), "
        "using the NIST Standard Reference Material 3600 (SRM 3600) glassy "
        "carbon as the primary calibrant (NIST, 2016). The software comprises a "
        "modular core library, a command-line interface (CLI), and a graphical "
        "user interface (GUI) with bilingual support (Chinese/English)."
    )

    add_para(doc,
        "The core library implements monitor-mode-aware normalization, robust "
        "K-factor estimation via median absolute deviation (MAD) outlier "
        "rejection, format-agnostic 1D profile parsing, and heterogeneous "
        "header extraction. Building on pyFAI (Ashiotis et al., 2015) and fabio, "
        "saxsabs adds the calibration-control and metadata-plumbing layers "
        "typically handled by ad hoc local scripts at synchrotron beamlines."
    )

    # ================================================================
    # STATEMENT OF NEED
    # ================================================================
    add_heading(doc, "Statement of need", level=1)

    add_para(doc,
        "Converting SAXS detector images to absolute-scale intensities requires "
        "dark-current subtraction, beam-monitor normalization, transmission and "
        "thickness correction, azimuthal integration, calibration against a "
        "reference standard, and batch reporting. In practice, each step is "
        "complicated by real-world heterogeneity: header formats differ across "
        "beamlines; metadata may reside in file headers, CSV tables, or manual "
        "input; and 1D profiles use inconsistent delimiters and column names."
    )

    add_para(doc,
        "Existing tools address individual stages well—pyFAI (Ashiotis et al., "
        "2015) for integration, SasView (Doucet et al., 2018) for model "
        "fitting, Dioptas (Prescher & Prakapenka, 2015) for 2D reduction, "
        "BioXTAS RAW (Hopkins et al., 2017) for biological SAXS, and Irena "
        "(Ilavsky & Jemian, 2009) for SAS modeling—but none provides a "
        "dedicated end-to-end absolute-calibration workflow that jointly handles "
        "metadata heterogeneity, multi-background averaging, robust K-factor "
        "estimation, and batch processing with audit trails."
    )

    add_para(doc,
        "saxsabs fills this gap, targeting beamline scientists and SAXS users "
        "who need reproducible, automatable absolute-scaling in production "
        "environments where metadata conventions are fluid and thousands of "
        "exposures per session are routine."
    )

    # ================================================================
    # STATE OF THE FIELD
    # ================================================================
    add_heading(doc, "State of the field", level=1)

    add_para(doc,
        "The SAXS ecosystem provides mature tools for individual processing "
        "stages: pyFAI for GPU-accelerated integration; SasView for model "
        "fitting; Dioptas for interactive 2D reduction; BioXTAS RAW for "
        "biological SAXS; DAWN (Basham et al., 2015) for plugin-based "
        "diffraction processing; and Irena for broad SAS analysis under "
        "Igor Pro."
    )

    add_para(doc,
        "Absolute intensity calibration remains a procedural gap. While the "
        "theory of calibration against NIST SRM 3600 glassy carbon is well "
        "documented (Glatter & Kratky, 1982; NIST, 2016), the operational "
        "workflow—parsing heterogeneous metadata, selecting normalization "
        "modes, handling multi-background subtraction, computing a robust "
        "scaling factor, and organizing traceable outputs—is typically left to "
        "bespoke scripts that are neither tested nor version-controlled."
    )

    add_para(doc, "Table 1 summarizes the functional landscape:", bold=False, italic=True)

    # -- Table 1 --
    capabilities = [
        ("Capability", "pyFAI", "SasView", "Dioptas", "BioXTAS RAW", "Irena", "saxsabs"),
        ("Azimuthal integration",          "✓", " ", "✓", "✓", "✓", " "),
        ("SAS model fitting",              " ", "✓", " ", "✓", "✓", " "),
        ("Heterogeneous header parsing",   " ", " ", " ", " ", " ", "✓"),
        ("Monitor-mode normalization",     " ", " ", " ", " ", " ", "✓"),
        ("Robust K-factor (MAD filtering)","" , " ", " ", " ", " ", "✓"),
        ("Format-agnostic 1D ingestion",   " ", " ", " ", "partial", " ", "✓"),
        ("Multi-background averaging",     " ", " ", " ", " ", " ", "✓"),
        ("Headless CLI + CI-testable",     "✓", "partial", " ", " ", " ", "✓"),
    ]

    table = doc.add_table(rows=len(capabilities), cols=7)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row_data in enumerate(capabilities):
        for c, cell_text in enumerate(row_data):
            cell = table.cell(r, c)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Times New Roman"
            if r == 0:  # header row
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    cap1 = doc.add_paragraph()
    cap1run = cap1.add_run("Table 1. Functional comparison of SAXS software tools. "
        "saxsabs focuses on the calibration-control and metadata-plumbing layer "
        "that bridges integration engines and absolute-scale reduction.")
    cap1run.italic = True
    cap1run.font.size = Pt(9)
    cap1run.font.name = "Times New Roman"
    cap1.paragraph_format.space_after = Pt(10)

    add_para(doc,
        "saxsabs complements these tools by formalizing the calibration-control "
        "layer that typically exists as private, untested scripts, making "
        "absolute-scaling workflows reproducible and auditable."
    )

    # ================================================================
    # SOFTWARE DESIGN
    # ================================================================
    add_heading(doc, "Software design", level=1)

    add_para(doc,
        "The primary design goal is to separate numerical calibration logic "
        "from UI concerns, enabling both interactive GUI use and headless "
        "CLI execution. The software is organized into four layers:"
    )

    add_rich_para(doc, [
        ("Core numerical layer", True, False),
        (" (saxsabs.core): implements monitor normalization and "
         "robust K-factor estimation as deterministic, stateless functions "
         "with no GUI or I/O side effects.", False, False),
    ])

    add_rich_para(doc, [
        ("I/O layer", True, False),
        (" (saxsabs.io): provides format-agnostic header parsing (fuzzy "
         "key matching with unit conversion) and multi-strategy 1D profile "
         "parsing (three separator strategies with automatic column-role "
         "inference by keyword matching).", False, False),
    ])

    add_rich_para(doc, [
        ("CLI layer", True, False),
        (" (saxsabs.cli): exposes four subcommands (norm-factor, "
         "parse-header, parse-external1d, estimate-k) for reproducible, "
         "scriptable execution.", False, False),
    ])

    add_rich_para(doc, [
        ("GUI layer", True, False),
        (" (SASAbs.py): a tkinter-based application with four tabbed "
         "panels—K-Factor Calibration, Batch Processing, External 1D "
         "Conversion, and Help—offering full bilingual (Chinese/English) "
         "user interface support.", False, False),
    ])

    add_figure(doc, "fig_workflow.png",
               "Figure 1. Calibration workflow of saxsabs. Input data (2D images, "
               "instrument metadata, and the NIST SRM 3600 reference) flow through "
               "header parsing, normalization, 2D background subtraction, pyFAI "
               "integration, and robust K-factor estimation to produce calibrated "
               "1D profiles with structured audit outputs.",
               width=Inches(5.8))

    add_para(doc,
        "This architecture enables unit testing of core functions independently "
        "of the GUI (15 automated tests across 3 OS × 3 Python versions), "
        "while preserving the GUI for interactive beamline use (Figure 3). The "
        "incremental migration from monolithic script to modular library avoids "
        "abrupt workflow disruption."
    )

    add_figure(doc, "fig_gui.png",
               "Figure 3. The saxsabs graphical user interface in English mode, "
               "showing the four-tab layout: K-Factor Calibration, Batch "
               "Processing, External 1D Conversion, and Help.",
               width=Inches(5.8))

    # -- Mathematical formulation --
    add_heading(doc, "Mathematical formulation", level=2)

    add_para(doc,
        "The absolute intensity calibration workflow in saxsabs follows the "
        "standard procedure documented for NIST SRM 3600 (NIST, 2016; "
        "Glatter & Kratky, 1982). The key computational steps are:"
    )

    # Normalization
    add_rich_para(doc, [
        ("Monitor normalization. ", True, False),
        ("Two modes are supported depending on whether the detector records "
         "count rates or integrated counts. In ", False, False),
        ("rate", False, True),
        (" mode, the normalization factor is:", False, False),
    ])
    add_equation(doc, "N = t_exp × I₀ × T", "1")
    add_para(doc,
        "where t_exp is the exposure time (s), I₀ is the beam-monitor count, "
        "and T is the sample transmission. In integrated mode, t_exp is omitted "
        "since the detector signal already accumulates over the full acquisition "
        "window:"
    )
    add_equation(doc, "N = I₀ × T", "2")

    # 2D background subtraction
    add_rich_para(doc, [
        ("2D background subtraction. ", True, False),
        ("Given sample detector image D_s, dark-current image D_d, and one or "
         "more background images {D_bg,i}, the net 2D scattering pattern is:", False, False),
    ])
    add_equation(doc,
        "I_net(x,y) = (D_s − D_d) / N_s − ⟨(D_bg,i − D_d) / N_bg,i⟩",
        "3"
    )
    add_para(doc,
        "where ⟨·⟩ denotes pixel-wise averaging (nanmean) over all available "
        "background images. This multi-background averaging reduces statistical "
        "noise in the background estimate."
    )

    # K-factor
    add_rich_para(doc, [
        ("Robust K-factor estimation. ", True, False),
        ("After azimuthal integration of the net pattern via pyFAI to obtain "
         "a 1D profile I_meas(q), the profile is interpolated onto the NIST "
         "SRM 3600 reference grid (15 data points spanning q ∈ [0.008, 0.250] "
         "Å⁻¹). Point-wise ratios are computed:", False, False),
    ])
    add_equation(doc, "R_i = I_ref(q_i) / I_meas(q_i)", "4")
    add_para(doc,
        "Outlier rejection uses the median absolute deviation (MAD):"
    )
    add_equation(doc, "σ̂ = 1.4826 × median(|R_i − R̃|)", "5")
    add_para(doc,
        "where R̃ = median(R_i). Points satisfying |R_i − R̃| > 3σ̂ are rejected, "
        "and the K-factor is the median of the remaining inlier ratios:"
    )
    add_equation(doc, "K = median(R_i)  for |R_i − R̃| ≤ 3σ̂", "6")
    add_para(doc,
        "The factor 1.4826 ensures consistency with the standard deviation "
        "under a Gaussian distribution. This robust estimator is resistant to "
        "outliers caused by parasitic scattering, beamstop shadows, or detector "
        "artefacts at the edges of the q-overlap region (Figure 2)."
    )

    add_figure(doc, "fig_kfactor_demo.png",
               "Figure 2. Demonstration of the robust K-factor estimation algorithm. "
               "(a) NIST SRM 3600 reference profile and a simulated measured profile "
               "after rescaling by K. (b) Point-wise ratios R_i = I_ref / I_meas with "
               "inlier points (green circles) and rejected outliers (red crosses); "
               "the blue line and shaded band show the median K-factor and ±3σ̂ "
               "acceptance region.",
               width=Inches(5.8))

    # Absolute conversion
    add_rich_para(doc, [
        ("Absolute intensity conversion. ", True, False),
        ("For each sample, the calibrated absolute intensity is:", False, False),
    ])
    add_equation(doc, "I_abs(q) = K × I_1D(q) / d", "7")
    add_para(doc,
        "where d is the sample thickness in centimeters. When transmission "
        "is available, the thickness can be estimated from the Beer–Lambert "
        "relation:"
    )
    add_equation(doc, "d = −ln(T) / μ", "8")
    add_para(doc,
        "where μ is the linear attenuation coefficient. For alloys or "
        "multi-element samples, μ is computed from the XCOM mass attenuation "
        "coefficients at the working energy:"
    )
    add_equation(doc, "μ = ρ × Σ(w_i × (μ/ρ)_i)", "9")
    add_para(doc,
        "where ρ is the bulk density, w_i is the mass fraction of element i, "
        "and (μ/ρ)_i is the mass attenuation coefficient."
    )

    # -- Batch processing features --
    add_heading(doc, "Batch processing and automation", level=2)

    add_para(doc,
        "The GUI batch-processing pipeline automates the complete chain from "
        "raw 2D images to calibrated 1D profiles. Key automation features "
        "include:"
    )

    bullets = [
        "Automatic background and dark-current matching via a weighted scoring "
        "function that compares exposure time, monitor counts, transmission, "
        "and temporal proximity between sample and candidate reference files.",

        "Multi-background capillary subtraction, where multiple background "
        "images are averaged pixel-wise to reduce statistical noise.",

        "Three azimuthal integration modes: full-ring, angular-sector (with "
        "±180° wrapping support), and radial chi-profile extraction.",

        "Sector merging with inverse-variance weighting: I = Σ(I_k × w_k) / Σ(w_k).",

        "Data quality controls including ≥98% non-positive-value detection "
        "and background normalization magnitude checking.",

        "Structured output traceability: each batch run produces a CSV report, "
        "a JSON metadata file, and a K-factor history log with timestamps and "
        "instrument parameters.",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(b)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

    # ================================================================
    # RESEARCH IMPACT STATEMENT
    # ================================================================
    add_heading(doc, "Research impact statement", level=1)

    add_para(doc,
        "saxsabs has been deployed for routine absolute intensity calibration "
        "at the Institute of Metal Research, Chinese Academy of Sciences, "
        "processing data from multiple synchrotron beamlines. It has replaced "
        "manual spreadsheet-based procedures, reducing operator intervention "
        "and eliminating errors from inconsistent header parsing."
    )

    add_para(doc,
        "The software defines its impact along three measurable dimensions:"
    )

    impact_items = [
        ("Operational efficiency: ", False,
         "Calibration previously requiring manual metadata extraction and "
         "iterative K-factor fitting is now a single CLI invocation or GUI "
         "session, reducing processing time from minutes to seconds."),

        ("Reliability: ", False,
         "Defensive parsing and format-agnostic ingestion have eliminated "
         "silent data-misinterpretation failures when switching between "
         "instruments."),

        ("Traceability: ", False,
         "Every run produces structured, deterministic output suitable for "
         "version control and audit."),
    ]

    for label, _, detail in impact_items:
        p = doc.add_paragraph(style="List Bullet")
        rl = p.add_run(label)
        rl.bold = True
        rl.font.size = Pt(11)
        rl.font.name = "Times New Roman"
        rd = p.add_run(detail)
        rd.font.size = Pt(11)
        rd.font.name = "Times New Roman"

    add_para(doc,
        "Core algorithms are verified by 15 automated tests across "
        "three operating systems and three Python versions under CI."
    )

    # ================================================================
    # AI USAGE DISCLOSURE
    # ================================================================
    add_heading(doc, "AI usage disclosure", level=1)

    add_para(doc,
        "The following AI-assisted coding tools were used during the "
        "development of this software:"
    )

    ai_bullets = [
        "GitHub Copilot (VS Code) and Anthropic Claude were used for code "
        "refactoring, internationalization extraction, test skeleton generation, "
        "and initial documentation drafts.",

        "AI assistance was limited to scaffolding and boilerplate. All core "
        "numerical algorithms were designed and validated by the author "
        "independently.",

        "Every AI-generated fragment was reviewed, tested, and revised before "
        "inclusion. Automated tests provide ongoing verification of scientific "
        "correctness.",
    ]
    for b in ai_bullets:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(b)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

    # ================================================================
    # ACKNOWLEDGEMENTS
    # ================================================================
    add_heading(doc, "Acknowledgements", level=1)

    add_para(doc,
        "The author thanks beamline scientists and users at the Institute of "
        "Metal Research who provided practical feedback on data heterogeneity "
        "and workflow failure modes during the development and deployment of "
        "this software."
    )

    # ================================================================
    # REFERENCES
    # ================================================================
    add_heading(doc, "References", level=1)

    references = [
        'Ashiotis, G., Deschiber, A., Nawber, M., Wright, J. P., Karkoulis, D., '
        'Picca, F. E., & Kieffer, J. (2015). The fast azimuthal integration '
        'Python library: pyFAI. Journal of Applied Crystallography, 48(2), '
        '510–519. https://doi.org/10.1107/S1600576715004306',

        'Basham, M., Filik, J., Wharmby, M. T., Chang, P. C. Y., El Sherif, B., '
        'Sheratt, R., ... & Hart, M. L. (2015). Data Analysis WorkbeNch (DAWN). '
        'Journal of Synchrotron Radiation, 22(3), 853–858. '
        'https://doi.org/10.1107/S1600577515002283',

        'Doucet, M., Cho, J. H., Alina, G., Bakber, J., Bouwman, W., Butler, P., '
        '... & Washington, A. (2018). SasView version 4.2. Zenodo. '
        'https://doi.org/10.5281/zenodo.1412041',

        'Glatter, O., & Kratky, O. (1982). Small Angle X-ray Scattering. '
        'Academic Press. ISBN 0-12-286280-5.',

        'Hopkins, J. B., Gillilan, R. E., & Ez, S. (2017). BioXTAS RAW: '
        'improvements to a free open-source program for small-angle X-ray '
        'scattering data reduction and analysis. Journal of Applied '
        'Crystallography, 50(5), 1545–1553. '
        'https://doi.org/10.1107/S1600576717011438',

        'Ilavsky, J., & Jemian, P. R. (2009). Irena: tool suite for modeling '
        'and analysis of small-angle scattering. Journal of Applied '
        'Crystallography, 42(2), 347–353. '
        'https://doi.org/10.1107/S0021889809002222',

        'National Institute of Standards and Technology. (2016). Standard '
        'Reference Material 3600: Absolute Intensity Calibration Standard for '
        'Small-Angle X-ray Scattering. Certificate of Analysis. '
        'https://www.nist.gov/srm',

        'Prescher, C., & Prakapenka, V. B. (2015). DIOPTAS: a program for '
        'reduction of two-dimensional X-ray diffraction data and data '
        'exploration. High Pressure Research, 35(3), 223–230. '
        'https://doi.org/10.1080/08957959.2015.1059835',
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        run = p.add_run(f"[{i}] {ref}")
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        p.paragraph_format.space_after = Pt(4)

    return doc


# ── entry point ────────────────────────────────────────────────────
if __name__ == "__main__":
    out_dir = Path(__file__).resolve().parent / "paper"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / "saxsabs_joss_paper.docx"

    doc = build_paper()
    doc.save(str(out_path))
    print(f"JOSS paper saved to: {out_path}")
